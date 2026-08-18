from __future__ import annotations


from hashlib import sha256

from io import BytesIO

from pathlib import Path

import re

from typing import (
    Literal,
)


from docx import (
    Document,
)

from pydantic import (
    BaseModel,
)

from pypdf import (
    PdfReader,
)


# ============================================================
# VERSION
# ============================================================

DOCUMENT_INGESTION_RULE_VERSION = (
    "document_ingestion_v0.2"
)


# ============================================================
# LIMITS
# ============================================================

SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
}


MAX_DOCUMENT_FILES = 20


MAX_DOCUMENT_BYTES = (
    25
    *
    1024
    *
    1024
)


DEFAULT_CHUNK_SIZE = 1200


DEFAULT_CHUNK_OVERLAP = 180


# ============================================================
# SCHEMAS
# ============================================================

class DocumentChunk(
    BaseModel
):
    chunk_id: str

    document_id: str

    filename: str

    extension: str

    chunk_index: int

    page_number: (
        int
        | None
    ) = None

    source_locator: str

    text: str

    character_count: int


class DocumentManifest(
    BaseModel
):
    document_id: str

    filename: str

    extension: str

    byte_count: int

    page_count: (
        int
        | None
    ) = None

    extracted_character_count: int

    chunk_count: int

    warnings: list[
        str
    ]


class DocumentIngestionReport(
    BaseModel
):
    status: Literal[
        "ready"
    ] = "ready"

    document_count: int

    chunk_count: int

    total_extracted_characters: int

    documents: list[
        DocumentManifest
    ]

    chunks: list[
        DocumentChunk
    ]

    warnings: list[
        str
    ]

    ingestion_rule_version: str = (
        DOCUMENT_INGESTION_RULE_VERSION
    )


# ============================================================
# INTERNAL SOURCE BLOCK
# ============================================================

class ExtractedTextBlock(
    BaseModel
):
    text: str

    page_number: (
        int
        | None
    ) = None

    source_locator: str


# ============================================================
# TEXT NORMALIZATION
# ============================================================

def normalize_text(
    value: str,
) -> str:
    value = (
        value
        .replace(
            "\r\n",
            "\n",
        )
        .replace(
            "\r",
            "\n",
        )
    )


    # Collapse repeated horizontal whitespace while preserving
    # line boundaries. This is particularly useful for pypdf's
    # layout extraction mode, which may insert several spaces
    # to represent the visual position of text.
    value = re.sub(
        r"[ \t]+",
        " ",
        value,
    )


    # Remove indentation introduced by page-layout extraction.
    value = re.sub(
        r"\n[ \t]+",
        "\n",
        value,
    )


    # Keep paragraph boundaries without allowing excessively
    # large vertical gaps.
    value = re.sub(
        r"\n{3,}",
        "\n\n",
        value,
    )


    return (
        value
        .strip()
    )


# ============================================================
# TEXT DECODING
# ============================================================

def decode_text_document(
    content: bytes,
) -> tuple[
    str,
    list[
        str
    ],
]:
    warnings: list[
        str
    ] = []


    try:
        return (
            content.decode(
                "utf-8-sig"
            ),
            warnings,
        )


    except UnicodeDecodeError:
        warnings.append(
            (
                "Le document texte n'était pas "
                "encodé en UTF-8. DataLens a "
                "utilisé un décodage Windows-1252."
            )
        )


        return (
            content.decode(
                "cp1252",
                errors="replace",
            ),
            warnings,
        )


# ============================================================
# PDF PAGE EXTRACTION
# ============================================================

def extract_pdf_page_text(
    page,
    *,
    page_index: int,
    warnings: list[
        str
    ],
) -> str:
    """
    Extract one PDF page.

    DataLens prefers pypdf's layout mode because it generally
    preserves word boundaries and list structure better for
    visually formatted business documents.

    If layout extraction fails or produces no usable text,
    DataLens falls back to the classic extraction mode.
    """

    layout_text = ""


    try:
        layout_text = (
            page.extract_text(
                extraction_mode=
                    "layout"
            )
            or ""
        )


    except Exception:
        warnings.append(
            (
                f"La page {page_index} n'a pas pu "
                "être extraite en mode layout. "
                "DataLens utilise le mode PDF "
                "classique pour cette page."
            )
        )


    normalized_layout = (
        normalize_text(
            layout_text
        )
    )


    if normalized_layout:
        return normalized_layout


    plain_text = ""


    try:
        plain_text = (
            page.extract_text()
            or ""
        )


    except Exception:
        warnings.append(
            (
                f"La page {page_index} "
                "n'a pas pu être extraite."
            )
        )

        return ""


    normalized_plain = (
        normalize_text(
            plain_text
        )
    )


    if normalized_plain:
        warnings.append(
            (
                f"La page {page_index} a été "
                "extraite avec le mode PDF "
                "classique car le mode layout "
                "n'a retourné aucun texte "
                "exploitable."
            )
        )


    return normalized_plain


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_blocks(
    content: bytes,
) -> tuple[
    list[
        ExtractedTextBlock
    ],
    int,
    list[
        str
    ],
]:
    warnings: list[
        str
    ] = []


    try:
        reader = PdfReader(
            BytesIO(
                content
            ),
            strict=False,
        )


    except Exception as error:
        raise ValueError(
            (
                "Le fichier PDF ne peut pas "
                "être lu par DataLens."
            )
        ) from error


    blocks: list[
        ExtractedTextBlock
    ] = []


    for (
        page_index,
        page,
    ) in enumerate(
        reader.pages,
        start=1,
    ):
        text = (
            extract_pdf_page_text(
                page,
                page_index=
                    page_index,
                warnings=
                    warnings,
            )
        )


        if not text:
            warnings.append(
                (
                    f"La page {page_index} "
                    "ne contient pas de texte "
                    "extractible."
                )
            )

            continue


        blocks.append(
            ExtractedTextBlock(
                text=
                    text,

                page_number=
                    page_index,

                source_locator=
                    f"page {page_index}",
            )
        )


    if not blocks:
        warnings.append(
            (
                "Aucun texte exploitable n'a "
                "été extrait du PDF. Le document "
                "est peut-être constitué "
                "principalement d'images."
            )
        )


    return (
        blocks,
        len(
            reader.pages
        ),
        warnings,
    )


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_blocks(
    content: bytes,
) -> tuple[
    list[
        ExtractedTextBlock
    ],
    list[
        str
    ],
]:
    warnings: list[
        str
    ] = []


    try:
        document = Document(
            BytesIO(
                content
            )
        )


    except Exception as error:
        raise ValueError(
            (
                "Le fichier DOCX ne peut pas "
                "être lu par DataLens."
            )
        ) from error


    blocks: list[
        ExtractedTextBlock
    ] = []


    paragraph_index = 0


    for paragraph in (
        document.paragraphs
    ):
        text = normalize_text(
            paragraph.text
        )


        if not text:
            continue


        paragraph_index += 1


        blocks.append(
            ExtractedTextBlock(
                text=
                    text,

                source_locator=(
                    "paragraphe "
                    f"{paragraph_index}"
                ),
            )
        )


    table_index = 0


    for table in (
        document.tables
    ):
        table_index += 1


        rows: list[
            str
        ] = []


        for row in table.rows:
            values = [
                normalize_text(
                    cell.text
                )

                for cell
                in row.cells
            ]


            values = [
                value

                for value
                in values

                if value
            ]


            if values:
                rows.append(
                    " | ".join(
                        values
                    )
                )


        if rows:
            blocks.append(
                ExtractedTextBlock(
                    text=
                        "\n".join(
                            rows
                        ),

                    source_locator=(
                        "tableau "
                        f"{table_index}"
                    ),
                )
            )


    if not blocks:
        warnings.append(
            (
                "Aucun texte exploitable n'a "
                "été extrait du document DOCX."
            )
        )


    return (
        blocks,
        warnings,
    )


# ============================================================
# TXT / MARKDOWN EXTRACTION
# ============================================================

def extract_text_blocks(
    content: bytes,
) -> tuple[
    list[
        ExtractedTextBlock
    ],
    list[
        str
    ],
]:
    (
        decoded,
        warnings,
    ) = decode_text_document(
        content
    )


    text = normalize_text(
        decoded
    )


    if not text:
        return (
            [],
            [
                *warnings,
                (
                    "Le document ne contient "
                    "pas de texte exploitable."
                ),
            ],
        )


    return (
        [
            ExtractedTextBlock(
                text=
                    text,

                source_locator=
                    "document",
            )
        ],
        warnings,
    )


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_document_blocks(
    *,
    filename: str,
    content: bytes,
) -> tuple[
    list[
        ExtractedTextBlock
    ],
    int | None,
    list[
        str
    ],
]:
    extension = (
        Path(
            filename
        )
        .suffix
        .lower()
    )


    if (
        extension
        not in SUPPORTED_DOCUMENT_EXTENSIONS
    ):
        raise ValueError(
            (
                "Format documentaire non pris "
                "en charge. Formats acceptés : "
                "PDF, DOCX, TXT et MD."
            )
        )


    if (
        extension
        ==
        ".pdf"
    ):
        (
            blocks,
            page_count,
            warnings,
        ) = extract_pdf_blocks(
            content
        )


        return (
            blocks,
            page_count,
            warnings,
        )


    if (
        extension
        ==
        ".docx"
    ):
        (
            blocks,
            warnings,
        ) = extract_docx_blocks(
            content
        )


        return (
            blocks,
            None,
            warnings,
        )


    (
        blocks,
        warnings,
    ) = extract_text_blocks(
        content
    )


    return (
        blocks,
        None,
        warnings,
    )


# ============================================================
# CHUNKING
# ============================================================

def split_text_into_chunks(
    text: str,
    *,
    chunk_size: int = (
        DEFAULT_CHUNK_SIZE
    ),
    overlap: int = (
        DEFAULT_CHUNK_OVERLAP
    ),
) -> list[
    str
]:
    text = normalize_text(
        text
    )


    if not text:
        return []


    if (
        chunk_size
        <=
        0
    ):
        raise ValueError(
            "chunk_size doit être positif."
        )


    if (
        overlap
        < 0
        or
        overlap
        >=
        chunk_size
    ):
        raise ValueError(
            (
                "overlap doit être positif "
                "ou nul et inférieur à "
                "chunk_size."
            )
        )


    if (
        len(
            text
        )
        <=
        chunk_size
    ):
        return [
            text
        ]


    chunks: list[
        str
    ] = []


    start = 0


    text_length = len(
        text
    )


    while (
        start
        <
        text_length
    ):
        tentative_end = min(
            start
            +
            chunk_size,

            text_length,
        )


        end = tentative_end


        if (
            tentative_end
            <
            text_length
        ):
            candidate = text[
                start:
                tentative_end
            ]


            boundary_positions = [
                candidate.rfind(
                    "\n\n"
                ),

                candidate.rfind(
                    ". "
                ),

                candidate.rfind(
                    " "
                ),
            ]


            boundary = max(
                boundary_positions
            )


            minimum_boundary = (
                chunk_size
                //
                2
            )


            if (
                boundary
                >=
                minimum_boundary
            ):
                if (
                    candidate[
                        boundary:
                        boundary + 2
                    ]
                    ==
                    ". "
                ):
                    end = (
                        start
                        +
                        boundary
                        +
                        1
                    )


                else:
                    end = (
                        start
                        +
                        boundary
                    )


        chunk = normalize_text(
            text[
                start:
                end
            ]
        )


        if chunk:
            chunks.append(
                chunk
            )


        if (
            end
            >=
            text_length
        ):
            break


        next_start = (
            end
            -
            overlap
        )


        if (
            next_start
            <=
            start
        ):
            next_start = (
                start
                +
                1
            )


        start = next_start


    return chunks


# ============================================================
# SINGLE DOCUMENT PROCESSING
# ============================================================

def process_document(
    *,
    filename: str,
    content: bytes,
) -> tuple[
    DocumentManifest,
    list[
        DocumentChunk
    ],
]:
    filename = (
        filename
        .strip()
    )


    if not filename:
        raise ValueError(
            (
                "Le document doit posséder "
                "un nom de fichier."
            )
        )


    extension = (
        Path(
            filename
        )
        .suffix
        .lower()
    )


    if (
        extension
        not in SUPPORTED_DOCUMENT_EXTENSIONS
    ):
        raise ValueError(
            (
                f"{filename} : format non "
                "pris en charge."
            )
        )


    if not content:
        raise ValueError(
            (
                f"{filename} : le fichier "
                "est vide."
            )
        )


    if (
        len(
            content
        )
        >
        MAX_DOCUMENT_BYTES
    ):
        raise ValueError(
            (
                f"{filename} dépasse la "
                "taille maximale autorisée "
                "pour cette version."
            )
        )


    document_hash = (
        sha256(
            content
        )
        .hexdigest()
    )


    document_id = (
        "document:"
        +
        document_hash[
            :16
        ]
    )


    (
        blocks,
        page_count,
        warnings,
    ) = extract_document_blocks(
        filename=
            filename,

        content=
            content,
    )


    chunks: list[
        DocumentChunk
    ] = []


    global_chunk_index = 0


    for block in blocks:
        block_chunks = (
            split_text_into_chunks(
                block.text
            )
        )


        for chunk_text in (
            block_chunks
        ):
            global_chunk_index += 1


            chunk_id = (
                f"{document_id}:chunk:"
                f"{global_chunk_index:04d}"
            )


            chunks.append(
                DocumentChunk(
                    chunk_id=
                        chunk_id,

                    document_id=
                        document_id,

                    filename=
                        filename,

                    extension=
                        extension,

                    chunk_index=
                        global_chunk_index,

                    page_number=
                        block.page_number,

                    source_locator=
                        block.source_locator,

                    text=
                        chunk_text,

                    character_count=
                        len(
                            chunk_text
                        ),
                )
            )


    extracted_character_count = sum(
        len(
            block.text
        )

        for block
        in blocks
    )


    manifest = DocumentManifest(
        document_id=
            document_id,

        filename=
            filename,

        extension=
            extension,

        byte_count=
            len(
                content
            ),

        page_count=
            page_count,

        extracted_character_count=
            extracted_character_count,

        chunk_count=
            len(
                chunks
            ),

        warnings=
            warnings,
    )


    return (
        manifest,
        chunks,
    )


# ============================================================
# MULTI-DOCUMENT INGESTION
# ============================================================

def build_document_ingestion_report(
    *,
    documents: list[
        tuple[
            str,
            bytes,
        ]
    ],
) -> DocumentIngestionReport:
    if not documents:
        raise ValueError(
            (
                "Au moins un document doit "
                "être fourni."
            )
        )


    if (
        len(
            documents
        )
        >
        MAX_DOCUMENT_FILES
    ):
        raise ValueError(
            (
                "Trop de documents ont été "
                "fournis en une seule fois."
            )
        )


    manifests: list[
        DocumentManifest
    ] = []


    all_chunks: list[
        DocumentChunk
    ] = []


    report_warnings: list[
        str
    ] = []


    seen_document_ids: set[
        str
    ] = set()


    for (
        filename,
        content,
    ) in documents:
        (
            manifest,
            chunks,
        ) = process_document(
            filename=
                filename,

            content=
                content,
        )


        if (
            manifest.document_id
            in seen_document_ids
        ):
            report_warnings.append(
                (
                    f"{filename} est un "
                    "doublon exact d'un "
                    "document déjà fourni "
                    "et n'a pas été indexé "
                    "une seconde fois."
                )
            )

            continue


        seen_document_ids.add(
            manifest.document_id
        )


        manifests.append(
            manifest
        )


        all_chunks.extend(
            chunks
        )


    return DocumentIngestionReport(
        document_count=
            len(
                manifests
            ),

        chunk_count=
            len(
                all_chunks
            ),

        total_extracted_characters=
            sum(
                manifest
                .extracted_character_count

                for manifest
                in manifests
            ),

        documents=
            manifests,

        chunks=
            all_chunks,

        warnings=
            report_warnings,
    )